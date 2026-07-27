"""生成 GB/T 9704-2012《党政机关公文格式》Word 公文模板。

文种清单、文件名前缀由 ``office_agent.domain.templates`` 的注册表决定，
本脚本只负责每个文种的【正文范例 spec】和渲染。两边文种对不上会直接报错，
不会悄悄生成孤儿文件。新增文种的完整步骤见《新增office模版的说明.md》。

每个模板含：版头（发文机关标志/发文字号/签发人）+ 红色分隔线 + 标题 +
主送机关 + 范例正文 + 落款 + 版记（抄送/印发）+ 页码。

用法::

    # 全量重新生成
    uv run --no-project --with python-docx python scripts/gen_official_templates.py
    # 只生成一个文种（改版式时快速迭代）
    uv run --no-project --with python-docx python scripts/gen_official_templates.py --only 通知
    # 只自检注册表与模板文件是否一致（不写文件，不需要 python-docx）
    python scripts/gen_official_templates.py --check

输出：template/word/01-决议.docx ... 15-纪要.docx
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

# 工程根：<root>/scripts/gen_official_templates.py
ROOT = Path(__file__).resolve().parent.parent
# 复用工程内的文种注册表；templates 模块只依赖标准库，
# 因此本脚本在 `--no-project`（未安装工程依赖）环境下也能 import 它。
sys.path.insert(0, str(ROOT / "src"))

from office_agent.domain.templates import (  # noqa: E402
    DOC_TYPE_NAMES,
    check_registry,
    template_path,
)

# --check 只读注册表和文件名，不渲染 docx；缺 python-docx 时推迟到真正生成时再报错。
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ModuleNotFoundError as e:  # pragma: no cover - 取决于运行环境
    _DOCX_MISSING: str | None = str(e)
else:
    _DOCX_MISSING = None

# ── 字体常量（GB/T 9704）───────────────────────────────────────────────
FANGSONG = "仿宋"  # 正文 / 主送 / 署名 / 版记
HEITI = "黑体"  # 一级标题（"一、"）
KAITI = "楷体"  # 二级标题（"（一）"）
XIAOBIAOSONG = "方正小标宋简体"  # 标题 / 版头发文机关标志（无该字体时自动回退到宋体）
EN_FONT = "Times New Roman"

# 页面可用宽度 = 21 − 2.8 − 2.6 = 15.6 cm（A4，GB/T 9704 页边距）
USABLE_WIDTH_CM = 15.6


# ── 基础工具 ────────────────────────────────────────────────────────────
def set_run_font(run, cn_font=FANGSONG, size_pt=None, bold=None, color=None):
    """设置 run 的中英文字体、字号、粗体、颜色。"""
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)


def style_para(p, line_pt=28, indent_chars=0, align=None, space_before=0, space_after=0):
    """段落级：固定行距、首行缩进、对齐、段前段后。"""
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_pt)
    if indent_chars:
        # 仿宋 16pt 一字，首行缩进按字号换算
        pf.first_line_indent = Pt(16 * indent_chars)
    if align is not None:
        p.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def add_para_border(p, position, sz="6", color="000000"):
    """给段落加一条边框线。position: top/bottom；sz 单位 1/8 pt。

    按 OOXML schema，``w:pBdr`` 必须位于 ``w:pPr`` 子元素顺序中的特定位置
    （``w:numPr`` 之后、``w:shd`` 之前），否则 schema 校验失败。
    """
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        # pBdr 合法前驱标签（按 CT_PPr 顺序，pBdr 紧随其后）
        predecessors = [
            "w:pStyle",
            "w:keepNext",
            "w:keepLines",
            "w:pageBreakBefore",
            "w:framePr",
            "w:widowControl",
            "w:numPr",
            "w:suppressLineNumbers",
        ]
        anchor = None
        for child in pPr:
            tag = child.tag
            if tag == qn("w:pBdr"):
                anchor = child
                break
            if tag.split("}", 1)[-1] in predecessors:
                anchor = child
        if anchor is not None:
            anchor.addnext(pBdr)
        else:
            pPr.insert(0, pBdr)
    b = OxmlElement(f"w:{position}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), sz)
    b.set(qn("w:space"), "0")
    b.set(qn("w:color"), color)
    pBdr.append(b)


def add_field(run, instr):
    """在 run 内插入域代码（如 PAGE）。"""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_el)
    run._r.append(fld_end)


# ── 公文版式组件 ────────────────────────────────────────────────────────
def set_page(section):
    """A4 + GB/T 9704 页边距（上3.7/下3.5/左2.8/右2.6 cm）。"""
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)


def add_org_header(doc, org_name):
    """版头发文机关标志（红色，小标宋，居中加大字号）。"""
    p = doc.add_paragraph()
    style_para(p, line_pt=48, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=12)
    r = p.add_run(org_name)
    set_run_font(r, XIAOBIAOSONG, size_pt=36, bold=True, color=(0xFF, 0x00, 0x00))


def add_doc_no_line(doc, doc_no, signer=None):
    """发文字号 + 签发人（上行文才显示签发人）。"""
    p = doc.add_paragraph()
    style_para(p, line_pt=28, space_before=0, space_after=0)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(USABLE_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(doc_no)
    set_run_font(r1, FANGSONG, size_pt=16)
    if signer:
        r2 = p.add_run("\t签发人：" + signer)
        set_run_font(r2, FANGSONG, size_pt=16)


def add_red_separator(doc):
    """版头与正文之间的红色分隔线（约 1.5pt）。"""
    p = doc.add_paragraph()
    style_para(p, line_pt=1, space_before=0, space_after=0)
    add_para_border(p, "bottom", sz="18", color="FF0000")


def add_title(doc, title):
    """公文标题（小标宋二号 22pt，居中，可折行）。"""
    p = doc.add_paragraph()
    style_para(p, line_pt=36, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=18, space_after=12)
    r = p.add_run(title)
    set_run_font(r, XIAOBIAOSONG, size_pt=22, bold=True)


def add_addressee(doc, addressee):
    """主送机关（仿宋三号，顶格）。"""
    p = doc.add_paragraph()
    style_para(p, line_pt=28, align=WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run(addressee)
    set_run_font(r, FANGSONG, size_pt=16)


def add_body_paragraph(doc, item):
    """根据 item 描述渲染正文段落（h1/h2/h3/p/closing/blank）。"""
    t = item.get("t", "p")
    text = item.get("text", "")

    if t == "blank":
        p = doc.add_paragraph()
        style_para(p, line_pt=28)
        return

    p = doc.add_paragraph()
    if t == "h1":  # 一级标题：黑体三号
        style_para(p, line_pt=28, indent_chars=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        r = p.add_run(text)
        set_run_font(r, HEITI, size_pt=16)
    elif t == "h2":  # 二级标题：楷体三号
        style_para(p, line_pt=28, indent_chars=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        r = p.add_run(text)
        set_run_font(r, KAITI, size_pt=16)
    elif t == "h3":  # 三级标题：仿宋三号加粗
        style_para(p, line_pt=28, indent_chars=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        r = p.add_run(text)
        set_run_font(r, FANGSONG, size_pt=16, bold=True)
    elif t == "attach":  # 附件说明：左对齐、不缩进
        style_para(p, line_pt=28, align=WD_ALIGN_PARAGRAPH.LEFT)
        r = p.add_run(text)
        set_run_font(r, FANGSONG, size_pt=16)
    elif t == "closing":  # 结语（独立段、缩进）
        style_para(p, line_pt=28, indent_chars=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        r = p.add_run(text)
        set_run_font(r, FANGSONG, size_pt=16)
    else:  # 正文 p：仿宋三号，首行缩进 2 字
        style_para(p, line_pt=28, indent_chars=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        r = p.add_run(text)
        set_run_font(r, FANGSONG, size_pt=16)


def add_signature(doc, org, date_cn):
    """落款：发文机关署名 + 成文日期（中文数字），右对齐。"""
    # 右空 4 字：用前导空格近似（公文规范右缩进 4 字）
    for line, space_before in ((org, 12), (date_cn, 0)):
        p = doc.add_paragraph()
        style_para(p, line_pt=28, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=space_before)
        r = p.add_run(line)
        set_run_font(r, FANGSONG, size_pt=16)


def add_record(doc, cc, issuer, issue_date):
    """版记：抄送栏 + 印发栏，黑色细线分隔。"""
    if not (cc or issuer):
        return
    # 抄送行（顶部细线）
    if cc:
        p_cc = doc.add_paragraph()
        style_para(p_cc, line_pt=28, space_before=12, space_after=0)
        add_para_border(p_cc, "top", sz="6", color="000000")
        r = p_cc.add_run("抄送：" + cc)
        set_run_font(r, FANGSONG, size_pt=16)
    # 印发行（上下细线）
    if issuer:
        p_iss = doc.add_paragraph()
        style_para(p_iss, line_pt=28, space_before=0, space_after=0)
        add_para_border(p_iss, "top", sz="6", color="000000")
        add_para_border(p_iss, "bottom", sz="6", color="000000")
        # 左右分布：制表位
        p_iss.paragraph_format.tab_stops.add_tab_stop(Cm(USABLE_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT)
        r1 = p_iss.add_run(issuer)
        set_run_font(r1, FANGSONG, size_pt=16)
        r2 = p_iss.add_run("\t" + issue_date)
        set_run_font(r2, FANGSONG, size_pt=16)


def add_page_number(section):
    """页脚页码：— X —（仿宋四号 14pt 居中）。"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run("— ")
    set_run_font(r1, FANGSONG, size_pt=14)
    r2 = p.add_run()
    set_run_font(r2, FANGSONG, size_pt=14)
    add_field(r2, r"PAGE \* MERGEFORMAT")
    r3 = p.add_run(" —")
    set_run_font(r3, FANGSONG, size_pt=14)


# ── 主装配 ──────────────────────────────────────────────────────────────
def build_doc(spec: dict) -> Document:
    doc = Document()
    # Normal 默认样式也改成仿宋，避免样式回退出现 Calibri
    normal = doc.styles["Normal"]
    normal.font.name = FANGSONG
    normal.font.size = Pt(16)
    normal_rPr = normal.element.get_or_add_rPr()
    rFonts = normal_rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        normal_rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FANGSONG)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)

    # 默认模板 lang.eastAsia=en-US（python-docx 自带 Normal.dotm），改为 zh-CN
    settings = doc.settings.element
    for lang in settings.findall(qn("w:lang")):
        if lang.get(qn("w:eastAsia")) in ("ja-JP", "en-US", None):
            lang.set(qn("w:eastAsia"), "zh-CN")
    # styles.xml 的 docDefaults 同步
    for lang in doc.styles.element.iter(qn("w:lang")):
        if lang.get(qn("w:eastAsia")) in ("ja-JP", "en-US", None):
            lang.set(qn("w:eastAsia"), "zh-CN")

    section = doc.sections[0]
    set_page(section)
    add_page_number(section)

    # —— 版头 ——
    if spec.get("org"):
        add_org_header(doc, spec["org"])
        add_doc_no_line(doc, spec["doc_no"], spec.get("signer"))
        add_red_separator(doc)

    # —— 标题 ——
    add_title(doc, spec["title"])

    # —— 主送机关 ——
    if spec.get("addressee"):
        add_addressee(doc, spec["addressee"])

    # —— 正文 ——
    for item in spec["body"]:
        add_body_paragraph(doc, item)

    # —— 落款 ——
    if spec.get("signer_org") or spec.get("date_cn"):
        # 落款前空一行
        add_body_paragraph(doc, {"t": "blank"})
        add_signature(doc, spec.get("signer_org", ""), spec.get("date_cn", ""))

    # —— 版记 ——
    add_record(doc, spec.get("cc", ""), spec.get("issuer", ""), spec.get("issue_date", ""))

    return doc


# ── 各文种的正文范例数据 ────────────────────────────────────────────────
# 每个文种一条 spec()，name 必须与 domain/templates.py 注册表里的文种名一致
# （顺序无所谓，输出文件名由注册表的 index 决定）。
# 模板默认元数据。
# 版头/落款/版记的固定槽位用 {{key}} 占位，供 officecli merge 一次性预填：
#   {{org}}        发文机关（版头红字 + 落款署名）
#   {{doc_no}}     发文字号（如 XX〔2026〕X号）
#   {{signer}}     签发人（仅上行文：请示/报告；其他文种不用）
#   {{date_cn}}    成文日期（中文数字，如 二〇二六年三月三十一日）
#   {{issuer}}     印发单位（版记印发栏左）
#   {{issue_date}} 印发日期（版记印发栏右）
# 正文范例里的 "XX市XX机关"、"XX工作" 等【不】改成占位——它们是供 LLM
# 编辑替换的示例文字，由 agent 用 set --find/--replace 或 set text= 处理。
ORG = "{{org}}"
DOC_NO = "{{doc_no}}"
ISSUER = "{{issuer}}"
ISSUE_DATE = "{{issue_date}}"
DATE_CN = "{{date_cn}}"
DATE_CN_SP = "二〇二六年三月三十一日"  # 演示规范写法（含“〇”），未使用


def spec(
    name,
    title,
    addressee,
    body,
    *,
    org=ORG,
    doc_no=DOC_NO,
    signer=None,
    signer_org="{{signer_org}}",
    date_cn=DATE_CN,
    cc="",
    issuer=ISSUER,
    issue_date=ISSUE_DATE,
):
    return dict(
        name=name,
        org=org,
        doc_no=doc_no,
        signer=signer,
        title=title,
        addressee=addressee,
        body=body,
        signer_org=signer_org,
        date_cn=date_cn,
        cc=cc,
        issuer=issuer,
        issue_date=issue_date,
    )


DOCUMENTS = [
    # ── 01 决议 ── 会议讨论通过的重大决策事项
    spec(
        "决议",
        "XX市XX机关第X届委员会第X次全体会议决议",
        "",
        [
            {"t": "p", "text": "（XX年X月X日XX市XX机关第X届委员会第X次全体会议通过）"},
            {
                "t": "p",
                "text": "XX市XX机关第X届委员会第X次全体会议，于XX年X月X日至X日在XX召开。出席这次全会的有市委委员X人，候补委员X人。有关方面负责同志列席了会议。",
            },
            {
                "t": "p",
                "text": "全会以习近平新时代中国特色社会主义思想为指导，全面贯彻党的二十大精神，听取和讨论了XX同志受市委常委会委托作的工作报告，审议通过了《中共XX市委关于XX的决定》。",
            },
            {"t": "h1", "text": "一、充分肯定XX工作取得的新成效"},
            {
                "t": "p",
                "text": "全会认为，市委X届X次全会以来，市委常委会团结带领全市广大干部群众，统筹推进各项事业，XX领域取得重要进展，为做好下一步工作打下了坚实基础。",
            },
            {"t": "h1", "text": "二、准确把握XX的新形势新任务"},
            {
                "t": "p",
                "text": "全会指出，当前XX工作正处于关键阶段，既面临重大机遇，也存在不少挑战。必须坚持稳中求进工作总基调，完整、准确、全面贯彻新发展理念，服务和融入新发展格局。",
            },
            {"t": "h1", "text": "三、扎实推进XX各项重点工作"},
            {"t": "h2", "text": "（一）着力推动高质量发展。"},
            {
                "t": "p",
                "text": "把发展经济的着力点放在实体经济上，加快构建现代化产业体系，推动XX、XX等优势产业提质增效。",
            },
            {"t": "h2", "text": "（二）持续深化改革扩大开放。"},
            {
                "t": "p",
                "text": "深入推进重点领域和关键环节改革，优化营商环境，激发各类经营主体活力，提升对外开放水平。",
            },
            {"t": "h2", "text": "（三）切实保障和改善民生。"},
            {
                "t": "p",
                "text": "用心用情办好就业、教育、医疗、养老等民生实事，不断增强人民群众获得感、幸福感、安全感。",
            },
            {"t": "h1", "text": "四、加强和改进党的领导"},
            {
                "t": "p",
                "text": "全会强调，做好XX工作，关键在党、关键在人。要坚定不移推进全面从严治党，营造风清气正的良好政治生态，凝聚干事创业的强大合力。",
            },
            {
                "t": "p",
                "text": "全会号召，全市各级党组织和广大党员干部要更加紧密地团结起来，埋头苦干、奋勇前进，为全面完成本次全会确定的各项任务而奋斗！",
            },
        ],
    ),
    # ── 02 决定 ── 重要事项作出决策和部署、奖惩
    spec(
        "决定",
        "XX市XX机关关于进一步加强XX工作的决定",
        "各区（县）XX机关，市属各有关部门：",
        [
            {
                "t": "p",
                "text": "为深入贯彻落实党中央、国务院关于XX工作的决策部署，推动我市XX事业高质量发展，现就进一步加强XX工作作出如下决定。",
            },
            {"t": "h1", "text": "一、充分认识加强XX工作的重要意义"},
            {
                "t": "p",
                "text": "XX是事关经济社会发展全局的重要工作。当前，我市XX工作基础不断夯实，但与高质量发展的要求相比，仍存在XX等突出问题。各区（县）、各部门要从战略和全局高度，切实增强责任感和紧迫感。",
            },
            {"t": "h1", "text": "二、明确XX工作的总体要求和主要目标"},
            {"t": "h2", "text": "（一）总体要求。"},
            {
                "t": "p",
                "text": "以习近平新时代中国特色社会主义思想为指导，坚持统筹谋划、突出重点、改革创新、依法推进，全面提升XX工作水平。",
            },
            {"t": "h2", "text": "（二）主要目标。"},
            {
                "t": "p",
                "text": "到XX年，XX体系更加健全，XX能力显著增强，XX指标达到XX水平，XX主要指标位居全省（区、市）前列。",
            },
            {"t": "h1", "text": "三、全面落实XX工作的重点任务"},
            {
                "t": "p",
                "text": "（一）完善XX体制机制。（二）强化XX要素保障。（三）推进XX重点项目建设。（四）深化XX领域改革。（五）加强XX人才队伍建设。",
            },
            {"t": "h1", "text": "四、切实加强组织领导"},
            {
                "t": "p",
                "text": "成立由市政府主要负责同志任组长的XX工作领导小组，统筹推进各项工作。各区（县）要落实属地责任，各有关部门要各司其职、密切配合，确保各项任务落到实处。",
            },
            {"t": "closing", "text": "XX工作的其他事项，由市XX局负责解释。"},
        ],
    ),
    # ── 03 命令（令）── 公布行政法规和规章、宣布施行重大强制性措施
    spec(
        "命令（令）",
        "XX市XX机关令",
        "",
        [
            {"t": "h2", "text": "第 X 号"},
            {
                "t": "p",
                "text": "《XX市XX办法》已经XX年X月X日XX市XX机关第X次常务会议通过，现予公布，自XX年X月X日起施行。",
            },
            {"t": "blank"},
            {"t": "p", "text": "机关长  XX"},
            {"t": "p", "text": "{{date_cn}}"},
            {"t": "blank"},
            {"t": "h2", "text": "XX市XX办法"},
            {
                "t": "p",
                "text": "第一条  为了XX，根据《XX法》和有关法律法规，结合本市实际，制定本办法。",
            },
            {"t": "p", "text": "第二条  本市行政区域内XX及其相关管理活动，适用本办法。"},
            {"t": "p", "text": "第三条  XX工作坚持XX原则，实行XX管理。"},
            {"t": "p", "text": "第四条  市、区（县）XX主管部门负责本行政区域内的XX监督管理工作。"},
            {"t": "p", "text": "第五条  本办法自XX年X月X日起施行。"},
        ],
        doc_no="{{doc_no}}",
        date_cn="{{date_cn}}",
    ),
    # ── 04 公报 ── 公布重要决定或重大事项
    spec(
        "公报",
        "XX市XX机关第X届委员会第X次全体会议公报",
        "",
        [
            {"t": "p", "text": "（XX年X月X日XX市XX机关第X届委员会第X次全体会议通过）"},
            {"t": "p", "text": "XX市XX机关第X届委员会第X次全体会议，于XX年X月X日至X日在XX举行。"},
            {
                "t": "p",
                "text": "出席这次全会的有，市委委员X人，候补委员X人。市纪律检查委员会常务委员会委员和有关方面负责同志列席了会议。市第X次党代会代表中部分基层同志和专家学者也列席了会议。",
            },
            {
                "t": "p",
                "text": "全会由市委常委会主持。全会听取和讨论了XX同志受市委常委会委托作的工作报告，审议通过了《中共XX市委关于制定XX市国民经济和社会发展第X个五年规划和二〇X X年远景目标的建议》。XX同志就《建议（讨论稿）》向全会作了说明。",
            },
            {
                "t": "p",
                "text": "全会充分肯定市委X届X次全会以来市委常委会的工作。一致认为，面对严峻复杂的形势和艰巨繁重的任务，市委常委会团结带领全市广大党员干部群众，攻坚克难、开拓进取，各项事业取得新的重大进展。",
            },
            {
                "t": "p",
                "text": "全会高度评价决胜全面建成小康社会取得的决定性成就。“十三五”时期，我市XX、XX等主要指标顺利完成，为开启全面建设社会主义现代化新征程奠定了坚实基础。",
            },
            {
                "t": "p",
                "text": "全会深入分析了新发展阶段XX面临的新机遇新挑战，提出了到二〇X X年基本实现社会主义现代化的远景目标，以及“十四五”时期经济社会发展指导思想、基本原则和主要目标。",
            },
            {
                "t": "p",
                "text": "全会提出了“十四五”时期经济社会发展主要任务：坚持创新驱动发展，建设现代化产业体系，形成强大国内市场，全面深化改革，优化国土空间布局，繁荣发展文化事业，推动绿色发展，实行高水平对外开放，改善人民生活品质，统筹发展和安全。",
            },
            {
                "t": "p",
                "text": "全会强调，实现“十四五”规划和二〇X X年远景目标，必须坚持党的全面领导，充分调动一切积极因素，广泛团结一切可以团结的力量，形成推动发展的强大合力。",
            },
            {
                "t": "p",
                "text": "全会号召，全市广大党员干部群众要更加紧密地团结起来，锐意进取、埋头苦干，为全面建设社会主义现代化XX作出新的更大贡献！",
            },
        ],
        issuer="",  # 公报一般不印发
    ),
    # ── 05 公告 ── 向国内外宣布重要事项或法定事项
    spec(
        "公告",
        "XX市XX机关公告",
        "",
        [
            {"t": "p", "text": "根据《XX法》《XX条例》的有关规定，现将XX事项公告如下："},
            {"t": "h1", "text": "一、XX范围"},
            {"t": "p", "text": "自XX年X月X日起，对XX区域（具体范围见附件）实行XX管理。"},
            {"t": "h1", "text": "二、XX内容"},
            {"t": "p", "text": "（一）XX期间，禁止从事XX活动。"},
            {"t": "p", "text": "（二）XX单位应当按照XX要求，落实XX措施。"},
            {"t": "p", "text": "（三）违反本公告规定的，由有关部门依法依规予以处理。"},
            {"t": "h1", "text": "三、其他事项"},
            {
                "t": "p",
                "text": "本公告自发布之日起施行。社会各界和广大群众应严格遵守，并予以监督。",
            },
            {"t": "closing", "text": "特此公告。"},
            {"t": "attach", "text": "附件：XX范围示意图"},
        ],
        date_cn="二〇二六年X月X日",
    ),
    # ── 06 通告 ── 在一定范围内公布应当遵守或周知的事项
    spec(
        "通告",
        "XX市XX机关关于XX的通告",
        "",
        [
            {
                "t": "p",
                "text": "为进一步加强XX管理，维护XX秩序，根据《XX法》《XX条例》等有关规定，现就XX有关事项通告如下：",
            },
            {"t": "h1", "text": "一、XX时间"},
            {"t": "p", "text": "自XX年X月X日零时起至XX年X月X日二十四时止。"},
            {"t": "h1", "text": "二、XX范围"},
            {"t": "p", "text": "本市行政区域内XX路段、XX场所。"},
            {"t": "h1", "text": "三、XX要求"},
            {"t": "h2", "text": "（一）XX期间，所有XX应当在指定区域内活动，不得进入XX区域。"},
            {"t": "h2", "text": "（二）XX单位和个人应当严格遵守本通告规定，服从现场管理人员指挥。"},
            {
                "t": "h2",
                "text": "（三）对违反本通告的行为，由相关主管部门依法予以处罚；构成犯罪的，依法追究刑事责任。",
            },
            {"t": "closing", "text": "特此通告。"},
            {"t": "attach", "text": "附件：XX区域明细表"},
        ],
        date_cn="二〇二六年X月X日",
    ),
    # ── 07 意见 ── 对重要问题提出见解和处理办法
    spec(
        "意见",
        "XX市XX机关关于XX工作的实施意见",
        "各区（县）XX机关，市属各有关部门：",
        [
            {
                "t": "p",
                "text": "为深入贯彻落实党中央、国务院关于XX工作的决策部署，扎实推进我市XX工作，现提出如下实施意见。",
            },
            {"t": "h1", "text": "一、总体要求"},
            {"t": "h2", "text": "（一）指导思想。"},
            {
                "t": "p",
                "text": "以习近平新时代中国特色社会主义思想为指导，全面贯彻党的二十大精神，坚持XX工作方针，推动XX事业高质量发展。",
            },
            {"t": "h2", "text": "（二）基本原则。"},
            {"t": "p", "text": "——坚持党的领导。充分发挥党总揽全局、协调各方的领导核心作用。"},
            {
                "t": "p",
                "text": "——坚持以人民为中心。把XX作为出发点和落脚点，让XX成果更多更公平惠及全体人民。",
            },
            {"t": "p", "text": "——坚持改革创新。破除制约XX的体制机制障碍，激发XX内生动力。"},
            {"t": "h2", "text": "（三）主要目标。"},
            {
                "t": "p",
                "text": "到XX年，XX体系基本建立，XX能力明显提升，XX指标较XX年翻一番，XX位居全省（区、市）前列。",
            },
            {"t": "h1", "text": "二、重点任务"},
            {"t": "h2", "text": "（一）加快XX建设。"},
            {"t": "p", "text": "统筹推进XX项目，确保XX年X月前建成投用。"},
            {"t": "h2", "text": "（二）深化XX改革。"},
            {"t": "p", "text": "推进XX领域“放管服”改革，简化XX审批流程，压缩XX办理时限。"},
            {"t": "h2", "text": "（三）强化XX保障。"},
            {"t": "p", "text": "加大XX投入力度，市财政每年安排专项资金X亿元，重点支持XX。"},
            {"t": "h1", "text": "三、保障措施"},
            {"t": "h2", "text": "（一）加强组织领导。"},
            {
                "t": "p",
                "text": "各区（县）要把XX工作摆上重要议事日程，主要负责同志亲自抓、负总责。",
            },
            {"t": "h2", "text": "（二）健全工作机制。"},
            {"t": "p", "text": "建立XX工作联席会议制度，定期研究解决重大问题。"},
            {"t": "h2", "text": "（三）严格考核监督。"},
            {"t": "p", "text": "将XX工作纳入年度目标考核，对工作落实不力的，严肃追责问责。"},
        ],
    ),
    # ── 08 通知 ── 发布、传达要求下级执行；批转、转发
    spec(
        "通知",
        "XX市XX机关关于做好XX工作的通知",
        "各区（县）XX机关，市属各有关部门：",
        [
            {
                "t": "p",
                "text": "为深入贯彻落实XX精神，切实做好XX工作，经市政府同意，现将有关事项通知如下：",
            },
            {"t": "h1", "text": "一、提高思想认识"},
            {
                "t": "p",
                "text": "XX是党中央、国务院作出的重大决策部署，事关XX全局。各区（县）、各部门要高度重视，切实把思想和行动统一到上级部署要求上来。",
            },
            {"t": "h1", "text": "二、把握重点环节"},
            {"t": "h2", "text": "（一）开展XX摸底排查。"},
            {
                "t": "p",
                "text": "各区（县）要于XX年X月X日前，对本辖区XX情况进行全面排查，建立XX台账。",
            },
            {"t": "h2", "text": "（二）制定XX工作方案。"},
            {"t": "p", "text": "结合实际，明确XX目标、任务和措施，于XX年X月X日前报市XX局备案。"},
            {"t": "h2", "text": "（三）抓好XX工作落实。"},
            {"t": "p", "text": "按照“XX”的要求，逐项抓好落实，确保XX任务如期完成。"},
            {"t": "h1", "text": "三、强化工作保障"},
            {
                "t": "p",
                "text": "（一）加强组织领导，明确责任分工；（二）强化监督检查，及时报送情况；（三）严肃工作纪律，对XX不力的依规追责。",
            },
            {"t": "h1", "text": "四、其他事项"},
            {
                "t": "p",
                "text": "请各区（县）、各部门于每月X日前，向市XX局报送工作进展情况。联系人：XX，联系电话：XXXX-XXXXXXX。",
            },
            {"t": "closing", "text": "请认真贯彻执行。"},
            {"t": "attach", "text": "附件：1.XX情况统计表  2.XX工作台账"},
        ],
    ),
    # ── 09 通报 ── 表彰先进、批评错误、传达重要精神和告知情况
    spec(
        "通报",
        "XX市XX机关关于表彰XX工作先进集体和先进个人的通报",
        "各区（县）XX机关，市属各有关部门：",
        [
            {
                "t": "p",
                "text": "近年来，全市各级XX机关和广大干部职工坚持以习近平新时代中国特色社会主义思想为指导，深入贯彻党中央、国务院决策部署，在XX工作中恪尽职守、开拓创新，涌现出一大批先进集体和先进个人。",
            },
            {
                "t": "p",
                "text": "为表彰先进、树立典型，激励全市XX系统干部职工见贤思齐、争创一流，经研究，决定对在XX工作中作出突出贡献的XX区XX局等X个先进集体和XX等X名先进个人予以通报表彰。",
            },
            {
                "t": "p",
                "text": "希望受表彰的先进集体和先进个人珍惜荣誉、再接再厉，充分发挥模范带头作用，在XX工作中再立新功。全市各级XX机关和广大干部职工要以先进为榜样，学习他们XX、XX的崇高品格，学习他们XX、XX的优良作风。",
            },
            {
                "t": "p",
                "text": "各级XX机关要广泛宣传先进事迹，营造学习先进、争当先进的浓厚氛围，激励广大干部职工以更加饱满的热情、更加务实的作风投身XX事业，为推动我市XX高质量发展作出新的更大贡献。",
            },
            {"t": "attach", "text": "附件：1.XX工作先进集体名单  2.XX工作先进个人名单"},
        ],
    ),
    # ── 10 报告 ── 向上级汇报工作、反映情况、回复询问（上行文）
    spec(
        "报告",
        "XX市XX机关关于XX工作的报告",
        "XX省XX厅：",
        [
            {
                "t": "p",
                "text": "按照《XX厅关于XX的通知》（X〔2026〕X号）要求，我市认真组织开展了XX工作，现将有关情况报告如下：",
            },
            {"t": "h1", "text": "一、主要工作情况"},
            {"t": "h2", "text": "（一）迅速安排部署。"},
            {
                "t": "p",
                "text": "市委、市政府高度重视，主要负责同志亲自部署，于X月X日召开专题会议，制定工作方案，明确责任分工。",
            },
            {"t": "h2", "text": "（二）扎实推进落实。"},
            {
                "t": "p",
                "text": "各区（县）、各部门按照统一安排，全面开展XX工作，累计出动人员X人次，排查XX场所X处，整改问题X个。",
            },
            {"t": "h2", "text": "（三）注重宣传引导。"},
            {
                "t": "p",
                "text": "通过报纸、电视、网络等媒体广泛宣传XX，发放宣传资料X万份，营造良好社会氛围。",
            },
            {"t": "h1", "text": "二、存在的主要问题"},
            {
                "t": "p",
                "text": "（一）部分基层单位XX力量薄弱，XX能力有待提升；（二）个别单位工作进展不平衡，存在XX等薄弱环节；（三）XX制度机制还需进一步完善。",
            },
            {"t": "h1", "text": "三、下一步工作打算"},
            {
                "t": "p",
                "text": "（一）持续巩固XX成果，防止问题反弹；（二）加大XX保障力度，提升XX能力；（三）健全长效机制，推动XX工作常态化、制度化。",
            },
            {"t": "closing", "text": "特此报告。"},
        ],
        signer="{{signer}}",
    ),
    # ── 11 请示 ── 向上级请求指示、批准（上行文，须签发人）
    spec(
        "请示",
        "XX市XX机关关于XX的请示",
        "XX省XX厅：",
        [
            {"t": "p", "text": "为XX，根据《XX》有关规定，我市拟XX。现将有关事项请示如下："},
            {"t": "h1", "text": "一、XX的必要性"},
            {
                "t": "p",
                "text": "近年来，我市XX发展迅速，XX需求持续增长。目前，XX已难以满足实际需要，亟需通过XX予以解决，这对于XX具有重要意义。",
            },
            {"t": "h1", "text": "二、XX的主要内容"},
            {"t": "h2", "text": "（一）XX范围。拟在XX区域实施XX，总规模XX。"},
            {
                "t": "h2",
                "text": "（二）资金安排。预计总投资X亿元，其中申请上级补助X亿元，市级配套X亿元，自筹X亿元。",
            },
            {"t": "h2", "text": "（三）实施时间。计划于XX年X月开工，XX年X月完工。"},
            {"t": "h1", "text": "三、请示事项"},
            {
                "t": "p",
                "text": "（一）恳请将XX纳入XX年度计划；（二）恳请在XX资金上予以支持X亿元；（三）恳请在XX政策上给予指导。",
            },
            {"t": "closing", "text": "妥否，请批示。"},
            {"t": "attach", "text": "附件：1.XX可行性研究报告  2.XX资金预算表"},
        ],
        signer="{{signer}}",
    ),
    # ── 12 批复 ── 答复下级请示事项（下行文）
    spec(
        "批复",
        "XX市XX机关关于XX的批复",
        "XX区XX机关：",
        [
            {"t": "p", "text": "你区《关于XX的请示》（X〔2026〕X号）收悉。经研究，现批复如下："},
            {"t": "h1", "text": "一、原则同意XX。"},
            {"t": "p", "text": "你区提出的XX方案符合实际，可按程序组织实施。"},
            {"t": "h1", "text": "二、关于XX资金问题。"},
            {
                "t": "p",
                "text": "市级财政在XX年度预算中统筹安排X万元予以支持，不足部分由你区自筹解决。",
            },
            {"t": "h1", "text": "三、关于XX实施要求。"},
            {
                "t": "p",
                "text": "（一）严格按照XX规范组织实施，确保XX质量；（二）加强XX资金管理，专款专用，严禁挪用；（三）XX项目完成后，及时将XX情况报市XX局备案。",
            },
            {"t": "p", "text": "请你区切实加强组织领导，落实工作责任，确保XX任务如期完成。"},
            {"t": "closing", "text": "此复。"},
        ],
    ),
    # ── 13 议案 ── 政府按法律程序向同级人大或其常委会提请审议
    spec(
        "议案",
        "关于XX的议案",
        "XX市人民代表大会常务委员会：",
        [
            {"t": "p", "text": "根据《XX法》和有关规定，现就XX事项提出议案，请予审议："},
            {"t": "h1", "text": "一、XX的必要性"},
            {
                "t": "p",
                "text": "近年来，我市XX快速发展，XX规模不断扩大。但与此同时，XX工作中也出现了一些新情况新问题，亟需通过XX加以规范和引导，这对于促进XX、保障XX具有重要意义。",
            },
            {"t": "h1", "text": "二、XX的主要内容"},
            {"t": "h2", "text": "（一）关于XX。明确XX的范围、标准和管理要求。"},
            {"t": "h2", "text": "（二）关于XX。规范XX的程序和方式，保障XX合法权益。"},
            {"t": "h2", "text": "（三）关于XX。完善XX监督机制，明确XX责任。"},
            {"t": "h2", "text": "（四）关于法律责任。对XX违法行为，设定相应的法律责任。"},
            {"t": "h1", "text": "三、需要说明的问题"},
            {
                "t": "p",
                "text": "（一）关于XX草案的起草过程。（二）关于XX草案的主要内容。（三）关于XX草案的可行性。以上说明连同XX草案，请予审议。",
            },
        ],
        org="{{org}}",
        signer_org="{{signer_org}}",
    ),
    # ── 14 函 ── 不相隶属机关之间商洽、询问和答复
    spec(
        "函",
        "XX市XX机关关于商请XX的函",
        "XX市XX局：",
        [
            {"t": "p", "text": "为做好XX工作，现就XX事宜商请贵局予以支持。"},
            {"t": "h1", "text": "一、XX背景"},
            {
                "t": "p",
                "text": "根据《XX》要求，我市正在组织开展XX工作。鉴于该项工作涉及贵局XX职责，为统筹推进、形成合力，特商请贵局协同配合。",
            },
            {"t": "h1", "text": "二、商请事项"},
            {"t": "h2", "text": "（一）请贵局明确X名同志作为XX工作联系人，参与日常工作协调。"},
            {"t": "h2", "text": "（二）请贵局及时提供XX相关资料和数据，配合做好XX分析研判。"},
            {"t": "h2", "text": "（三）请贵局在XX事项上给予业务指导和政策支持。"},
            {"t": "h1", "text": "三、其他事项"},
            {
                "t": "p",
                "text": "我局已明确XX处XX同志为联系人（联系电话：XXXX-XXXXXXX），具体事宜可由双方对接落实。",
            },
            {"t": "closing", "text": "请予大力支持为盼。"},
        ],
    ),
    # ── 15 纪要 ── 记载会议主要情况和议定事项（一般不设版记/署名）
    spec(
        "纪要",
        "XX市XX机关XX年第X次XX会议纪要",
        "",
        [
            {
                "t": "p",
                "text": "XX年X月X日，XX同志主持召开XX市XX机关XX年第X次XX会议，研究XX、XX等事项。XX、XX等同志出席，XX、XX等同志列席。现将会议议定事项纪要如下。",
            },
            {"t": "h1", "text": "一、关于XX工作"},
            {
                "t": "p",
                "text": "会议听取了XX局关于XX工作情况的汇报。会议认为，今年以来，全市XX工作取得积极成效，XX、XX等指标稳步提升，但XX、XX等方面仍需加大力度。",
            },
            {
                "t": "p",
                "text": "会议议定：（一）原则同意XX局提出的XX工作方案，由XX局牵头组织实施。（二）由XX同志牵头，协调解决XX、XX等问题。（三）市财政安排X万元专项用于XX。",
            },
            {"t": "h1", "text": "二、关于XX事项"},
            {
                "t": "p",
                "text": "会议听取了XX局关于XX事项的汇报。会议指出，XX事关XX，要按照XX原则，统筹推进、稳妥实施。",
            },
            {
                "t": "p",
                "text": "会议议定：（一）同意XX方案，按程序报批后实施。（二）各区（县）、各有关部门要按照职责分工，密切配合，确保XX顺利推进。（三）XX局要加强跟踪调度，重大事项及时报告。",
            },
            {"t": "h1", "text": "三、关于XX工作"},
            {
                "t": "p",
                "text": "会议强调，当前XX工作正处于关键时期，各区（县）、各部门要切实增强责任感，压实工作责任，确保各项议定事项落到实处。",
            },
            {"t": "p", "text": "出席：XX、XX、XX"},
            {"t": "p", "text": "列席：XX、XX、XX"},
            {"t": "p", "text": "送：市委各部委，市人大常委会办公室，市政府各部门，市政协办公室"},
        ],
        signer_org="",
        date_cn="",
        issuer="",
        issue_date="",
    ),
]


# ── 主入口 ─────────────────────────────────────────────────────────────
SPEC_BY_NAME: dict[str, dict] = {d["name"]: d for d in DOCUMENTS}


def check_specs() -> list[str]:
    """校验本脚本的正文 spec 与注册表文种是否一一对应。"""
    problems = []
    registered = set(DOC_TYPE_NAMES)
    for name in sorted(registered - set(SPEC_BY_NAME)):
        problems.append(
            f"{name}: 注册表里有，但本脚本缺正文 spec"
            f"（在 DOCUMENTS 末尾加一条 spec('{name}', ...)）"
        )
    for name in sorted(set(SPEC_BY_NAME) - registered):
        problems.append(f"{name}: 本脚本有正文 spec，但没在 domain/templates.py 的 _DEFS 里注册")
    return problems


def generate(names: list[str]) -> None:
    """渲染指定文种的模板，输出路径由注册表决定。"""
    if _DOCX_MISSING:
        raise SystemExit(
            f"生成模板需要 python-docx（{_DOCX_MISSING}）。\n"
            f"请用: uv run --no-project --with python-docx python "
            f"scripts/gen_official_templates.py"
        )
    for name in names:
        path = template_path(name)
        path.parent.mkdir(parents=True, exist_ok=True)
        build_doc(SPEC_BY_NAME[name]).save(path)
        print(f"  ✓ {path.name}")
    print(f"完成，共生成 {len(names)} 个公文模板。")


def main(argv: list[str] | None = None) -> None:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument(
        "--only",
        metavar="文种",
        action="append",
        help="只生成指定文种（可重复），如 --only 通知 --only 请示",
    )
    parser.add_argument(
        "--check",
        action="store_true",
        help="只自检（注册表 ↔ 正文 spec ↔ 模板文件），不写文件",
    )
    args = parser.parse_args(argv)

    problems = check_specs()
    if args.check:
        problems += check_registry()
        if problems:
            print("自检发现问题：")
            for p in problems:
                print(f"  ✗ {p}")
            raise SystemExit(1)
        print(f"自检通过：{len(DOC_TYPE_NAMES)} 个文种，注册表 / 正文 spec / 模板文件一致。")
        return

    if problems:
        print("注册表与正文 spec 不一致，已中止生成：")
        for p in problems:
            print(f"  ✗ {p}")
        raise SystemExit(1)

    if args.only:
        unknown = [n for n in args.only if n not in SPEC_BY_NAME]
        if unknown:
            raise SystemExit(f"未知文种 {unknown}。合法文种: {DOC_TYPE_NAMES}")
        names = args.only
    else:
        names = list(DOC_TYPE_NAMES)

    print(f"输出目录：{template_path(names[0]).parent}")
    generate(names)


if __name__ == "__main__":
    main()
